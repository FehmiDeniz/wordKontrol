import os
from docx import Document
from docx.shared import Inches
from docx import Document
from docx.text.paragraph import Paragraph
def Iceriyomu(dosyayol):
    document = Document('{}'.format(dosyayol))
    headings = []
    texts = []
    para = []
    giris = ""
    for paragraph in document.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            if headings:
                texts.append(para)
            headings.append(paragraph.text)
            para = []
        elif paragraph.style.name == "Normal" and not paragraph.text.find(' ',0,1) != -1 and paragraph.text !='':
            para.append(paragraph.text)
    if para or len(headings)>len(texts):
        texts.append(texts.append(para))

    for h, t in zip(headings, texts):
        if h== "GİRİŞ" or h== "Giriş":
            giris = t[-1]
            #print(giris)



    if (giris.find('apsam') != -1 or giris.find('rganizasyon') != -1): 
        sonuc="Giris bölümünün son bölümünde tezin organizasyonu ve kapsamına yer verilmis "
        RaporaEkle(sonuc)
    else:
        sonuc="Giris bölümünün son bölümünde tezin organizasyonu ve kapsamına yer verilmemis"
        RaporaEkle(sonuc)

def RaporaEkle(sonuc):
    
    f = open('WordRapor.docx', 'rb')

    document = Document(f)

    
    document.add_paragraph(
        sonuc, style='List Number'
        )

    document.add_heading('16541504-Fatih Uludag', level=1)
    document.add_heading('175541058-Doğukan Kurnaz', level=1)
    document.add_heading('14545520-Kemal Sanlı', level=1)
    document.add_heading('175541059-Batuhan Harmanşah', level=1)
    document.save('WordRapor.docx')

    f.close()

    print("Asama uc tamamlandi...✓")
    print("Word Raporu Olusturuldu...✓")
