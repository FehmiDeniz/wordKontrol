import docx
import os
from docx import Document
from docx.shared import Inches

def cifttirnakfonk(dosyayol):
   
    doc = docx.Document('{}'.format(dosyayol))
    content = doc.paragraphs
    liste=""

    for p in content:
    
        liste=liste+p.text
    yeni=[]
    sayac=0

    #print(liste)
    while (sayac<len(liste)):
    

        if (liste[sayac] =="”" or liste[sayac] =="“"):
            while(sayac<len(liste)):
                sayac=sayac+1
                if(len(liste)==sayac):
                    break
                elif(liste[sayac]=="”" or liste[sayac] =="“"):
                    yeni.append("------")    
                    break
                else: 
                    yeni.append(liste[sayac])
        sayac=sayac+1
    kelimesay=0
    for i in yeni:
   
        if(i== " "):
            kelimesay=kelimesay+1
    if(kelimesay>50):
        sonuc="Cift tirnak arasinda kullanılan kelime adedi>50"
        RaporOlustur(sonuc)
    else:
        sonuc="Cift tirnak arasinda kullanılan kelime adedi>50 degil"
        RaporOlustur(sonuc)
def RaporOlustur(degisgen):
    document = Document()

    document.add_heading('WORD KONTROL RAPORU', 0)

    
    document.add_paragraph(
        degisgen, style='List Number'
    )


    

    document.save('WordRapor.docx')
    print("            ╔═╗┬─┐┌─┐┌─┐┬─┐┌─┐┌┬┐┬  ┌─┐┌┬┐┌─┐  ╔╦╗┬┬  ┬  ┌─┐┬─┐┬")
    print("            ╠═╝├┬┘│ ││ ┬├┬┘├─┤││││  ├─┤│││├─┤   ║║││  │  ├┤ ├┬┘│")
    print("            ╩  ┴└─└─┘└─┘┴└─┴ ┴┴ ┴┴─┘┴ ┴┴ ┴┴ ┴  ═╩╝┴┴─┘┴─┘└─┘┴└─┴")
        
    print("Asama Bir Tamamlandi...✓")
    print("islem Word Raporuna eklendi....✓")
    
